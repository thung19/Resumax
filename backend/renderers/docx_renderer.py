"""DOCX Renderer.

Reproduces the source document by walking the captured element sequence
and replaying each paragraph's exact formatting properties.

When content changes (tailoring), the renderer replaces text while
preserving the original formatting structure.
"""

from __future__ import annotations

from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from backend.models.resume_content import (
    ResumeContent,
    ResumeSection,
    SectionType,
)
from backend.models.resume_ir import ResumeIR
from backend.models.resume_layout import (
    ElementType,
    LayoutElement,
    ParagraphFormat,
    ResumeLayout,
    RunFormat,
    TabAlignment,
    TabStop,
)


class DocxRenderer:
    """Render a ResumeIR to a .docx file."""

    def __init__(self, ir: ResumeIR):
        self._ir = ir
        self._content = ir.content
        self._layout = ir.layout
        self._doc = Document()
        page = self._layout.page
        self._content_width_twips = int(
            (page.width_in - page.margin_left_in - page.margin_right_in) * 1440
        )

        # Use values extracted from the original document
        self._font = self._layout.body_font_family or self._detect_font()
        self._spacer_half_pt = self._layout.spacer_size_half_pt or self._detect_spacer_size()
        self._body_line_spacing_val = self._layout.body_line_spacing_val

    def _detect_spacer_size(self) -> int:
        """Fallback: pick the most common spacer size from elements."""
        from collections import Counter
        sizes: Counter = Counter()
        for el in self._layout.elements:
            if el.element_type == ElementType.SPACER and el.spacer_size_half_pt is not None:
                sizes[el.spacer_size_half_pt] += 1
        if sizes:
            return sizes.most_common(1)[0][0]
        return 10  # 5pt default

    def _detect_font(self) -> str:
        """Fallback: detect font from styles."""
        for role in ["bullet", "entry_header", "name", "contact"]:
            style = self._layout.styles.get(role)
            if style and style.font.family and style.font.family.lower() not in ("arial", "minorhansi"):
                return style.font.family
        return "Garamond"

    def render(self) -> bytes:
        self._setup_page()
        if self._layout.elements:
            self._render_from_elements()
        else:
            self._render_from_content()
        buf = BytesIO()
        self._doc.save(buf)
        return buf.getvalue()

    # --- Page setup ---

    def _setup_page(self):
        page = self._layout.page
        section = self._doc.sections[0]
        section.page_width = Inches(page.width_in)
        section.page_height = Inches(page.height_in)
        section.top_margin = Inches(page.margin_top_in)
        section.bottom_margin = Inches(page.margin_bottom_in)
        section.left_margin = Inches(page.margin_left_in)
        section.right_margin = Inches(page.margin_right_in)

        # Normal style: use values extracted from the original document
        ns = self._layout.normal_style
        normal = self._doc.styles["Normal"]
        normal.font.name = ns.font_family
        normal.font.size = Pt(ns.font_size_pt)

        pPr = normal.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            normal.element.append(pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:line"), str(ns.line_spacing_val))
        spacing.set(qn("w:lineRule"), ns.line_rule)
        spacing.set(qn("w:before"), str(ns.space_before_twips))
        spacing.set(qn("w:after"), str(ns.space_after_twips))

    # --- OOXML helpers ---

    def _apply_paragraph_format(self, para: Paragraph, pf: ParagraphFormat):
        """Apply captured paragraph formatting to a python-docx paragraph."""
        pPr = para._element.get_or_add_pPr()

        # Spacing – always write explicit before/after to prevent
        # Word Normal-style defaults from leaking extra space.
        has_spacing = (
            pf.line_value is not None
            or pf.space_before_twips is not None
            or pf.space_after_twips is not None
        )
        if has_spacing:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            if pf.line_value is not None:
                spacing.set(qn("w:line"), str(pf.line_value))
            if pf.line_rule is not None:
                spacing.set(qn("w:lineRule"), pf.line_rule)
            # Only write before/after if they were explicitly set
            if pf.space_before_twips is not None:
                spacing.set(qn("w:before"), str(pf.space_before_twips))
            if pf.space_after_twips is not None:
                spacing.set(qn("w:after"), str(pf.space_after_twips))

        # Indentation
        if any([pf.indent_left_twips, pf.indent_right_twips,
                pf.indent_hanging_twips, pf.indent_first_line_twips]):
            ind = pPr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                pPr.append(ind)
            if pf.indent_left_twips is not None:
                ind.set(qn("w:left"), str(pf.indent_left_twips))
            if pf.indent_right_twips is not None:
                ind.set(qn("w:right"), str(pf.indent_right_twips))
            if pf.indent_hanging_twips is not None:
                ind.set(qn("w:hanging"), str(pf.indent_hanging_twips))
            if pf.indent_first_line_twips is not None:
                ind.set(qn("w:firstLine"), str(pf.indent_first_line_twips))

        # Alignment
        if pf.alignment:
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                pPr.append(jc)
            jc.set(qn("w:val"), pf.alignment)

        # Tab stops
        if pf.tab_stops:
            tabs = pPr.find(qn("w:tabs"))
            if tabs is None:
                tabs = OxmlElement("w:tabs")
                pPr.append(tabs)
            for ts in pf.tab_stops:
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), ts.alignment.value)
                tab.set(qn("w:pos"), str(ts.position_twips))
                tabs.append(tab)

        # Bottom border
        if pf.bottom_border and pf.bottom_border.enabled:
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), str(int(pf.bottom_border.width_pt * 8)))
            bottom.set(qn("w:space"), "1")
            if pf.bottom_border.color:
                bottom.set(qn("w:color"), pf.bottom_border.color)
            else:
                bottom.set(qn("w:color"), "000000")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # Keep with next
        if pf.keep_with_next:
            kwn = OxmlElement("w:keepNext")
            pPr.append(kwn)

        # Word style
        if pf.word_style:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is None:
                pStyle = OxmlElement("w:pStyle")
                pPr.insert(0, pStyle)
            pStyle.set(qn("w:val"), pf.word_style)

    def _apply_run_format(self, run: Run, rf: RunFormat):
        """Apply captured run formatting."""
        rPr = run._element.get_or_add_rPr()

        # Font family
        fam = rf.font_family or self._font
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), fam)
        rFonts.set(qn("w:hAnsi"), fam)
        rFonts.set(qn("w:eastAsia"), fam)
        rFonts.set(qn("w:cs"), fam)

        # Size
        if rf.font_size_half_pt is not None:
            for tag in ["w:sz", "w:szCs"]:
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn("w:val"), str(rf.font_size_half_pt))

        # Bold
        if rf.bold is True:
            if rPr.find(qn("w:b")) is None:
                rPr.append(OxmlElement("w:b"))
        if rf.bold_cs is True:
            if rPr.find(qn("w:bCs")) is None:
                rPr.append(OxmlElement("w:bCs"))

        # Italic
        if rf.italic is True:
            if rPr.find(qn("w:i")) is None:
                rPr.append(OxmlElement("w:i"))
        if rf.italic_cs is True:
            if rPr.find(qn("w:iCs")) is None:
                rPr.append(OxmlElement("w:iCs"))

        # Underline
        if rf.underline:
            u = rPr.find(qn("w:u"))
            if u is None:
                u = OxmlElement("w:u")
                rPr.append(u)
            u.set(qn("w:val"), rf.underline)

        # Color
        if rf.color:
            c = rPr.find(qn("w:color"))
            if c is None:
                c = OxmlElement("w:color")
                rPr.append(c)
            c.set(qn("w:val"), rf.color)

        # Small caps
        if rf.small_caps:
            if rPr.find(qn("w:smallCaps")) is None:
                rPr.append(OxmlElement("w:smallCaps"))

    def _add_runs_to_paragraph(self, para: Paragraph, runs: list[RunFormat]):
        """Add runs with their exact formatting to a paragraph."""
        for rf in runs:
            if rf.hyperlink_url:
                # Create a w:hyperlink element with the run inside
                self._add_hyperlink_run(para, rf)
                continue

            run = para.add_run()
            if rf.is_tab:
                run._element.append(OxmlElement("w:tab"))
                if rf.text:
                    t = OxmlElement("w:t")
                    t.set(qn("xml:space"), "preserve")
                    t.text = rf.text
                    run._element.append(t)
            elif rf.text:
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = rf.text
                run._element.append(t)
            self._apply_run_format(run, rf)

    def _add_hyperlink_run(self, para: Paragraph, rf: RunFormat):
        """Add a hyperlink run to a paragraph using OOXML."""
        # Add relationship
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        part = self._doc.part
        r_id = part.relate_to(rf.hyperlink_url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Build hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Build run inside hyperlink
        run_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        # Hyperlink style
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        # Font
        if rf.font_family:
            rFonts = OxmlElement("w:rFonts")
            for a in ["ascii", "hAnsi", "eastAsia", "cs"]:
                rFonts.set(qn(f"w:{a}"), rf.font_family)
            rPr.append(rFonts)
        if rf.font_size_half_pt:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(rf.font_size_half_pt))
            rPr.append(sz)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), str(rf.font_size_half_pt))
            rPr.append(szCs)
        if rf.bold:
            rPr.append(OxmlElement("w:b"))
        run_el.append(rPr)

        # Text
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = rf.text
        run_el.append(t)

        hyperlink.append(run_el)
        para._element.append(hyperlink)

    def _set_paragraph_rPr(self, para: Paragraph, half_pt: int, font: str):
        """Set pPr/rPr — paragraph-level default run properties.

        Controls the visual height and font of empty (spacer) paragraphs.
        Both font size AND font family must be set, otherwise Word falls
        back to the Normal style font which may differ.
        """
        pPr = para._element.get_or_add_pPr()
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            pPr.append(rPr)
        # Font family
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rFonts.set(qn(f"w:{attr}"), font)
        # Font size
        for tag in ["w:sz", "w:szCs"]:
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), str(half_pt))

    def _clean_lr_paragraphs(self, pf: ParagraphFormat) -> list[ParagraphFormat]:
        """Collapse tab-separated paragraph(s) into clean left/right layouts.

        Source documents often use 5-15 consecutive tab runs to visually
        push dates to the right.  We collapse these into:
          [left text runs] [single tab] [right text runs]
        with a single right-aligned tab stop at the content width.

        Some source docs pack TWO logical lines (company+date AND
        role+location) into a single paragraph.  We detect this and
        split into separate paragraph formats.
        """
        import re

        has_tab_run = any(r.is_tab for r in pf.runs)
        # Also check for date-after-padding (no tabs but spaces + date run)
        has_date_padding = False
        if not has_tab_run and len(pf.runs) >= 2:
            for idx in range(1, len(pf.runs)):
                prev_t = pf.runs[idx - 1].text
                curr_t = pf.runs[idx].text.lstrip()
                if prev_t.endswith("  ") and curr_t and re.search(
                    r"^(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b",
                    curr_t, re.I,
                ):
                    has_date_padding = True
                    break
        if not has_tab_run and not has_date_padding:
            return [pf]

        # Split runs into groups separated by:
        # 1. Tab sequences
        # 2. Bold→non-bold transitions with leading whitespace (new logical line)
        # 3. Runs that start with a date pattern after whitespace-padded text
        groups: list[list[RunFormat]] = []
        current: list[RunFormat] = []
        prev_bold: bool | None = None
        for r in pf.runs:
            if r.is_tab:
                if current:
                    groups.append(current)
                    current = []
                prev_bold = None
            else:
                is_bold = r.bold is True
                text_stripped = r.text.lstrip()
                leading_spaces = len(r.text) - len(text_stripped)

                # Case 1: bold→non-bold with leading whitespace (role after date)
                bold_transition = (
                    current
                    and prev_bold is True
                    and not is_bold
                    and leading_spaces >= 3
                )

                # Case 2: date-like run after whitespace-padded previous run
                # (e.g., "Atlantic Health System     " + " Jun 2025 – Aug 2025")
                date_after_padding = False
                if (
                    current
                    and not bold_transition
                    and text_stripped
                    and re.search(
                        r"^(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b",
                        text_stripped, re.I,
                    )
                ):
                    prev_text = current[-1].text if current else ""
                    if prev_text.endswith("  "):  # 2+ trailing spaces
                        date_after_padding = True

                if bold_transition or date_after_padding:
                    # Split: trim trailing spaces from previous group
                    if current and current[-1].text:
                        current[-1] = current[-1].model_copy()
                        current[-1].text = current[-1].text.rstrip()
                    groups.append(current)
                    current = []

                current.append(r)
                prev_bold = is_bold
        if current:
            groups.append(current)

        # Filter out empty/whitespace-only groups
        groups = [
            g for g in groups
            if "".join(r.text for r in g).strip()
        ]

        if not groups:
            return [pf]

        # Classify each group: does it look like a date/location?
        DATE_RE = re.compile(
            r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|present)\b|\b\d{4}\b",
            re.I,
        )
        LOCATION_RE = re.compile(
            r",\s*[A-Z]{2}\b|^Remote$|^remote$",
        )

        def _is_date_or_location(text: str) -> bool:
            t = text.strip()
            if not t:
                return False
            return bool(DATE_RE.search(t) or LOCATION_RE.search(t))

        def _group_text(group: list[RunFormat]) -> str:
            return "".join(r.text for r in group).strip()

        # Pair up groups: [left, right, left, right, ...]
        # A "right" group is one that looks like a date or location.
        pairs: list[tuple[list[RunFormat], list[RunFormat]]] = []
        i = 0
        while i < len(groups):
            left = groups[i]
            right: list[RunFormat] = []
            # Check if next group is a date/location (right side)
            if i + 1 < len(groups) and _is_date_or_location(_group_text(groups[i + 1])):
                right = groups[i + 1]
                i += 2
            else:
                # Check if THIS group contains embedded right content
                # e.g., "Sep 2024 – Present                  Research Analyst"
                # which means the left+right are merged into one group
                left_text = _group_text(left)
                if _is_date_or_location(left_text) and not pairs:
                    # This group IS the right side of a missing left
                    # (shouldn't normally happen)
                    right = left
                    left = []
                    i += 1
                else:
                    i += 1
            pairs.append((left, right))

        # Build paragraph formats for each pair
        result: list[ParagraphFormat] = []
        for left, right in pairs:
            pf_copy = pf.model_copy(deep=True)

            # Clean excessive whitespace from left and right edges
            # but preserve single spaces between words (run boundaries)
            left_clean = [r.model_copy() for r in left]
            right_clean = [r.model_copy() for r in right]

            # Only strip leading whitespace from the very first run
            # and trailing whitespace from the very last run
            if left_clean:
                left_clean[0].text = left_clean[0].text.lstrip()
                left_clean[-1].text = left_clean[-1].text.rstrip()
            if right_clean:
                right_clean[0].text = right_clean[0].text.lstrip()
                right_clean[-1].text = right_clean[-1].text.rstrip()

            # Remove runs that are only whitespace (padding runs)
            left_clean = [r for r in left_clean if r.text]
            right_clean = [r for r in right_clean if r.text]

            left_text = "".join(r.text for r in left_clean).strip()
            right_text = "".join(r.text for r in right_clean).strip()

            if not left_text and not right_text:
                continue

            if right_text and _is_date_or_location(right_text):
                # Build left + tab + right
                tab_fmt = right_clean[0] if right_clean else (
                    left_clean[0] if left_clean else RunFormat()
                )
                tab_run = RunFormat(
                    text="",
                    is_tab=True,
                    font_family=tab_fmt.font_family,
                    font_size_half_pt=tab_fmt.font_size_half_pt,
                    bold=tab_fmt.bold,
                )
                pf_copy.runs = left_clean + [tab_run] + right_clean
                pf_copy.tab_stops = [TabStop(
                    position_twips=self._content_width_twips,
                    alignment=TabAlignment.RIGHT,
                )]
            else:
                # No right content — just left side
                pf_copy.runs = left_clean or right_clean
                pf_copy.tab_stops = []

            if pf_copy.runs:
                result.append(pf_copy)

        return result if result else [pf]

    def _replay_paragraph(self, pf: ParagraphFormat) -> Paragraph:
        """Create a paragraph that exactly replays captured formatting."""
        para = self._doc.add_paragraph()
        self._apply_paragraph_format(para, pf)
        self._add_runs_to_paragraph(para, pf.runs)

        # Set paragraph-level rPr (pPr/rPr) — use the original's
        # captured values if available.
        if pf.ppr_font_size_half_pt or pf.ppr_font_family:
            half_pt = pf.ppr_font_size_half_pt or 0
            font = pf.ppr_font_family or self._font
            pPr = para._element.get_or_add_pPr()
            rPr = pPr.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                pPr.append(rPr)
            # Always set font family if we have it
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                rFonts.set(qn(f"w:{attr}"), font)
            # Only set size if the original had one
            if half_pt:
                for tag in ["w:sz", "w:szCs"]:
                    el = rPr.find(qn(tag))
                    if el is None:
                        el = OxmlElement(tag)
                        rPr.append(el)
                    el.set(qn("w:val"), str(half_pt))

        return para

    # --- Element-based rendering ---

    def _render_from_elements(self):
        """Walk the element sequence and reproduce each paragraph."""
        # Remove default empty paragraph
        if self._doc.paragraphs:
            p = self._doc.paragraphs[0]._element
            p.getparent().remove(p)

        # Build a content lookup for substitutions
        content_map = self._build_content_map()

        # Track which content we've rendered so we can substitute text
        section_idx = 0
        entry_idx = 0
        bullet_idx = 0
        current_section = None

        for el in self._layout.elements:
            pf = el.paragraph_format

            if el.element_type == ElementType.SPACER:
                # Replay spacer as a clean empty paragraph.
                # Preserve original spacing exactly — don't add
                # attributes the original didn't have.
                pf_copy = pf.model_copy(deep=True)
                pf_copy.runs = []
                pf_copy.tab_stops = []
                pf_copy.word_style = None
                para = self._replay_paragraph(pf_copy)
                self._set_paragraph_rPr(
                    para,
                    el.spacer_size_half_pt or self._spacer_half_pt,
                    self._font,
                )

            elif el.element_type in (ElementType.NAME, ElementType.CONTACT):
                # Replay with potentially updated contact info
                self._replay_paragraph(pf)

            elif el.element_type == ElementType.SECTION_HEADING:
                # Replay heading — convert drawing-based line to border
                pf_copy = pf.model_copy(deep=True)
                if pf.has_drawing and not pf.bottom_border:
                    from backend.models.resume_layout import BorderDef
                    pf_copy.bottom_border = BorderDef(
                        enabled=True, width_pt=0.75, color="000000"
                    )
                # Remove the empty drawing run and any invisible-color runs
                clean_runs = []
                for r in pf_copy.runs:
                    if r.color and r.color.upper() in ("FFFFFF", "WHITE"):
                        continue
                    if not r.text.strip() and not r.is_tab and not r.font_size_half_pt:
                        continue
                    clean_runs.append(r)
                pf_copy.runs = clean_runs
                self._replay_paragraph(pf_copy)

            elif el.element_type == ElementType.ENTRY_HEADER:
                for cleaned_pf in self._clean_lr_paragraphs(pf):
                    self._replay_paragraph(cleaned_pf)

            elif el.element_type == ElementType.ENTRY_SUBHEADER:
                for cleaned_pf in self._clean_lr_paragraphs(pf):
                    self._replay_paragraph(cleaned_pf)

            elif el.element_type == ElementType.BULLET:
                self._replay_paragraph(pf)

            elif el.element_type == ElementType.SKILLS_ROW:
                self._replay_paragraph(pf)

            else:
                self._replay_paragraph(pf)

    def _build_content_map(self) -> dict:
        """Build a map for content substitution during tailoring."""
        return {}

    # --- Fallback content-based rendering (when no elements captured) ---

    def _render_from_content(self):
        """Fallback renderer when no element sequence is available."""
        if self._doc.paragraphs:
            p = self._doc.paragraphs[0]._element
            p.getparent().remove(p)

        contact = self._content.contact
        font = self._font

        if contact.name:
            para = self._doc.add_paragraph()
            self._set_spacing(para, 240, "auto")
            self._set_align(para, "center")
            run = para.add_run(contact.name)
            self._set_run_font(run, font, half_pt=52, bold=True)

        parts = []
        if contact.phone: parts.append(contact.phone)
        if contact.email: parts.append(contact.email)
        if contact.linkedin: parts.append(contact.linkedin)
        if contact.github: parts.append(contact.github)
        if parts:
            para = self._doc.add_paragraph()
            self._set_align(para, "center")
            run = para.add_run(" | ".join(parts))
            self._set_run_font(run, font)

        for section in self._content.sections:
            self._render_section_fallback(section)

    def _render_section_fallback(self, section: ResumeSection):
        """Fallback section renderer."""
        font = self._font

        # Spacer
        self._add_empty_para(font)

        if section.title:
            para = self._doc.add_paragraph()
            self._set_spacing(para, 240, "auto")
            self._add_border(para)
            run = para.add_run(section.title)
            self._set_run_font(run, font, bold=True)

        self._add_empty_para(font, half_pt=10)

        if section.type == SectionType.EDUCATION:
            for e in section.education_entries:
                date = e.end_date or ""
                self._add_lr_para(font, e.institution, date, bold=True)
                degree = e.degree or ""
                self._add_lr_para(font, degree, e.location or "")
                if e.coursework:
                    self._add_empty_para(font, half_pt=10)
                    para = self._doc.add_paragraph()
                    self._set_spacing(para, 240, "auto")
                    r1 = para.add_run("Coursework: ")
                    self._set_run_font(r1, font, bold=True)
                    r2 = para.add_run(", ".join(e.coursework))
                    self._set_run_font(r2, font)

        elif section.type in (SectionType.EXPERIENCE, SectionType.VOLUNTEER):
            for i, e in enumerate(section.experience_entries):
                if i > 0:
                    self._add_empty_para(font)
                date = ""
                if e.start_date and e.end_date:
                    date = f"{e.start_date} \u2013 {e.end_date}"
                elif e.start_date:
                    # Currently-employed entries have a start_date and
                    # no end_date \u2014 this used to render with no date at
                    # all (pdf_renderer.py and text_renderer.py both
                    # already fell back correctly here).
                    date = e.start_date
                elif e.end_date:
                    date = e.end_date
                self._add_lr_para(font, e.company, date, bold=True)
                self._add_lr_para(font, e.role, e.location or "", italic=True)
                for b in e.bullets:
                    para = self._doc.add_paragraph()
                    self._set_spacing(para, 240, "auto")
                    run = para.add_run(f"\u2022 {b.text}")
                    self._set_run_font(run, font)

        elif section.type == SectionType.PROJECTS:
            for i, e in enumerate(section.project_entries):
                if i > 0:
                    self._add_empty_para(font)
                # pdf_renderer.py/text_renderer.py both include the date
                # range when present \u2014 this fallback never read
                # start_date/end_date at all, so project dates were
                # silently absent from this render path.
                date = ""
                if e.start_date and e.end_date:
                    date = f"{e.start_date} \u2013 {e.end_date}"
                elif e.start_date:
                    date = e.start_date
                elif e.end_date:
                    date = e.end_date
                if date:
                    self._add_lr_para(font, e.name, date, bold=True)
                else:
                    para = self._doc.add_paragraph()
                    self._set_spacing(para, 240, "auto")
                    run = para.add_run(e.name)
                    self._set_run_font(run, font, bold=True)
                for b in e.bullets:
                    para = self._doc.add_paragraph()
                    self._set_spacing(para, 240, "auto")
                    run = para.add_run(f"\u2022 {b.text}")
                    self._set_run_font(run, font)

        elif section.type == SectionType.SKILLS:
            for cat in section.skill_categories:
                para = self._doc.add_paragraph()
                self._set_spacing(para, 240, "auto")
                r1 = para.add_run(f"{cat.category}: ")
                self._set_run_font(r1, font, bold=True)
                r2 = para.add_run(", ".join(cat.skills))
                self._set_run_font(r2, font)

        else:
            # Certifications/Awards/Publications/Custom sections parse
            # into GenericEntry, not one of the branches above \u2014 this
            # had no else at all, so such a section rendered as just its
            # bare heading with the entry (and any unstructured
            # raw_lines) entirely dropped.
            for i, e in enumerate(section.generic_entries):
                if i > 0:
                    self._add_empty_para(font)
                if e.title:
                    para = self._doc.add_paragraph()
                    self._set_spacing(para, 240, "auto")
                    run = para.add_run(e.title)
                    self._set_run_font(run, font, bold=True)
                if e.subtitle:
                    para = self._doc.add_paragraph()
                    self._set_spacing(para, 240, "auto")
                    run = para.add_run(e.subtitle)
                    self._set_run_font(run, font, italic=True)
                for b in e.bullets:
                    para = self._doc.add_paragraph()
                    self._set_spacing(para, 240, "auto")
                    run = para.add_run(f"\u2022 {b.text}")
                    self._set_run_font(run, font)

        for line in section.raw_lines:
            para = self._doc.add_paragraph()
            self._set_spacing(para, 240, "auto")
            run = para.add_run(line)
            self._set_run_font(run, font)

    # --- Simple helpers for fallback ---

    def _set_spacing(self, para, line, rule):
        pPr = para._element.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), rule)
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        pPr.append(sp)

    def _set_align(self, para, val):
        pPr = para._element.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), val)
        pPr.append(jc)

    def _set_run_font(self, run, family, half_pt=None, bold=False, italic=False):
        rPr = run._element.get_or_add_rPr()
        rf = OxmlElement("w:rFonts")
        for a in ["ascii", "hAnsi", "eastAsia", "cs"]:
            rf.set(qn(f"w:{a}"), family)
        rPr.insert(0, rf)
        if half_pt:
            for t in ["w:sz", "w:szCs"]:
                el = OxmlElement(t)
                el.set(qn("w:val"), str(half_pt))
                rPr.append(el)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))

    def _add_empty_para(self, font, half_pt=None):
        para = self._doc.add_paragraph()
        self._set_spacing(para, 240, "auto")
        run = para.add_run()
        self._set_run_font(run, font, half_pt=half_pt)

    def _add_border(self, para):
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_lr_para(self, font, left, right, bold=False, italic=False):
        para = self._doc.add_paragraph()
        self._set_spacing(para, 240, "auto")
        # Right tab stop
        pPr = para._element.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(self._content_width_twips))
        tabs.append(tab)
        pPr.append(tabs)

        r1 = para.add_run(left)
        self._set_run_font(r1, font, bold=bold, italic=italic)
        rt = para.add_run()
        self._set_run_font(rt, font, bold=bold, italic=italic)
        rt._element.append(OxmlElement("w:tab"))
        r2 = para.add_run(right)
        self._set_run_font(r2, font, bold=bold, italic=italic)
